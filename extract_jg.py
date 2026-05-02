import os
import json
from docx import Document
import re

doc_path = r"C:\Users\hashm\Desktop\Projects\Plan Mode\AH\book\Johar Gandhi Book.docx"
if not os.path.exists(doc_path):
    print("DOCX file not found!")
    exit(1)

doc = Document(doc_path)

book_data = {
  "book": {
    "title": "Johar Gandhi",
    "author": "Amir Hashmi",
    "regNo": "L-108073/2021",
    "releaseDate": "October 2, 2022"
  },
  "chapters": []
}

inventory = []

keywords = {
  "freedom_fighters": [],
  "tribes": [],
  "events": [],
  "locations": []
}

current_chapter = None
current_section = None
chapter_counter = 0

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    
    # Very basic heuristic for Chapter
    lower_text = text.lower()
    if "chapter" in lower_text and len(text) < 50:
        chapter_counter += 1
        current_chapter = {
            "id": chapter_counter,
            "title": text,
            "sections": [],
            "imagePlaceholders": []
        }
        book_data["chapters"].append(current_chapter)
        current_section = None
        continue
        
    if not current_chapter:
        # Before any chapter, maybe preface
        chapter_counter += 1
        current_chapter = {
            "id": chapter_counter,
            "title": "Preface",
            "sections": [],
            "imagePlaceholders": []
        }
        book_data["chapters"].append(current_chapter)
        
    # Heuristic for section heading: Short text, Title Case or all caps, no period at end
    is_heading = len(text) < 100 and (text.isupper() or text.istitle()) and not text.endswith('.')
    # Wait, some text might just be short paragraphs. We can also check para style
    if para.style.name.startswith('Heading') or is_heading:
        current_section = {
            "heading": text,
            "paragraphs": []
        }
        current_chapter["sections"].append(current_section)
    else:
        if not current_section:
            current_section = {
                "heading": "Introduction",
                "paragraphs": []
            }
            current_chapter["sections"].append(current_section)
            
        current_section["paragraphs"].append(text)
        
        # Add image placeholders roughly every 4 paragraphs
        if len(current_section["paragraphs"]) % 4 == 0:
            prompt = "Historical illustration relating to " + " ".join(text.split()[:10])
            img_path = f"/images/johar-gandhi/chapter{current_chapter['id']}/img_{len(current_chapter['imagePlaceholders'])+1:03d}.jpg"
            placeholder = {
                "position": f"after_paragraph_{len(current_section['paragraphs'])}",
                "prompt": prompt,
                "path": img_path
            }
            current_chapter["imagePlaceholders"].append(placeholder)
            inventory.append({
                "chapterId": current_chapter["id"],
                "path": img_path,
                "prompt": prompt,
                "status": "pending"
            })

# Extract some keywords blindly
all_text = " ".join([p.text for p in doc.paragraphs])
import string
words = [w.strip(string.punctuation) for w in all_text.split() if w[0].isupper() and len(w) > 4]
from collections import Counter
top_words = [w[0] for w in Counter(words).most_common(50)]
keywords["locations"] = top_words[:10]
keywords["events"] = top_words[10:20]
keywords["tribes"] = top_words[20:30]
keywords["freedom_fighters"] = top_words[30:40]

os.makedirs('data', exist_ok=True)
with open('data/book-content.json', 'w', encoding='utf-8') as f:
    json.dump(book_data, f, ensure_ascii=False, indent=2)

with open('data/keywords.json', 'w', encoding='utf-8') as f:
    json.dump(keywords, f, ensure_ascii=False, indent=2)

with open('data/image-inventory.json', 'w', encoding='utf-8') as f:
    json.dump(inventory, f, ensure_ascii=False, indent=2)

print("Extraction complete!")
